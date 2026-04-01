"""
nodes.py – LangGraph node implementations

OCR/text extraction  →  Gemini 1.5 Flash  (fast, cheap, excellent vision)
Grading              →  Claude claude-haiku-4-5 Sonnet (strong reasoning)
"""

import base64
import logging
import os
from pathlib import Path

import google.generativeai as genai
from anthropic import Anthropic
from docx import Document

from app.prompts import EXTRACT_TEXT_PROMPT, GRADE_SYSTEM_PROMPT, GRADE_USER_PROMPT
from app.state import GradingState

logger = logging.getLogger(__name__)

# ── Helpers ──────────────────────────────────────────────────────────────────


def _encode_image(image_path: str) -> tuple[str, str]:
    """Return (base64_data, mime_type) for the given image file."""
    path = Path(image_path)
    ext = path.suffix.lower()
    mime_map = {
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".png": "image/png",
        ".webp": "image/webp",
        ".gif": "image/gif",
    }
    mime_type = mime_map.get(ext, "image/jpeg")
    with open(image_path, "rb") as f:
        data = base64.b64encode(f.read()).decode("utf-8")
    return data, mime_type


def _read_docx_text(doc_path: str) -> str:
    """Extract plain text from a .docx file, preserving paragraph breaks."""
    doc = Document(doc_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)


# ── Node 1: Load teacher's model answer from .docx ───────────────────────────


def load_teacher_answers_node(state: GradingState) -> dict:
    """
    Reads the teacher's model-answer .docx and stores its text in state.
    """
    doc_path = state["teacher_doc_path"]
    logger.info("Loading teacher answers from: %s", doc_path)

    if not os.path.exists(doc_path):
        raise FileNotFoundError(
            f"Teacher answer document not found: {doc_path}\n"
            "Please create the file and re-run."
        )

    teacher_text = _read_docx_text(doc_path)
    if not teacher_text.strip():
        raise ValueError(
            f"Teacher answer document is empty: {doc_path}"
        )

    logger.info("Teacher answers loaded (%d chars).", len(teacher_text))
    return {"teacher_answers": teacher_text}


# ── Node 2: Extract text from student image using Gemini ─────────────────────


def extract_text_node(state: GradingState) -> dict:
    """
    Uses Gemini 1.5 Flash (vision) to OCR the student's handwritten image.
    """
    image_path = state["image_path"]
    logger.info("Extracting text from image: %s", image_path)

    if not os.path.exists(image_path):
        raise FileNotFoundError(f"Student image not found: {image_path}")

    # Configure Gemini
    genai.configure(api_key=os.environ["GEMINI_API_KEY"])
    model = genai.GenerativeModel("gemini-1.5-flash")

    image_data, mime_type = _encode_image(image_path)

    response = model.generate_content(
        [
            EXTRACT_TEXT_PROMPT,
            {"mime_type": mime_type, "data": image_data},
        ]
    )

    extracted = response.text.strip()
    if not extracted:
        raise ValueError("Gemini returned empty text — check the image quality.")

    logger.info("Text extracted (%d chars).", len(extracted))
    return {"extracted_text": extracted}


# ── Node 3: Save extracted text to a .docx file ──────────────────────────────


def save_doc_node(state: GradingState) -> dict:
    """
    Saves the student's extracted answer as a formatted .docx in data/outputs/.
    """
    doc = Document()
    doc.add_heading("Student Answer — Transcription", level=0)

    # Source image reference
    doc.add_heading("Source Image", level=2)
    doc.add_paragraph(state["image_path"])

    # Extracted content
    doc.add_heading("Extracted Text", level=2)
    for line in state["extracted_text"].split("\n"):
        doc.add_paragraph(line)

    os.makedirs("data/outputs", exist_ok=True)
    doc_path = "data/outputs/student_answer.docx"
    doc.save(doc_path)

    logger.info("Student answer saved to: %s", doc_path)
    return {"doc_path": doc_path}


# ── Node 4: Grade with Claude ─────────────────────────────────────────────────


def grade_node(state: GradingState) -> dict:
    """
    Uses Claude claude-haiku-4-5 Sonnet to grade the student's answer against the
    teacher's model answer and rubric.
    """
    logger.info("Grading with Claude…")

    client = Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    user_message = GRADE_USER_PROMPT.format(
        teacher_answers=state["teacher_answers"],
        rubric=state["rubric"],
        student_answer=state["extracted_text"],
    )

    response = client.messages.create(
        model="claude-haiku-4-5-20251001",   # Fast + cost-effective for grading
        max_tokens=1024,
        system=GRADE_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )

    grade_text = response.content[0].text.strip()
    logger.info("Grading complete.")
    return {"final_grade": grade_text}